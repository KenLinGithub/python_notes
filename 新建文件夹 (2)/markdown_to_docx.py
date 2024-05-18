from docx import Document

doc = Document()

with open('sample.txt', encoding='utf8') as f:
    lines = f.readlines()

para = ''
for line in lines:
    line = line.strip()

    if line == '':
        if para != '':
            doc.add_paragraph(para)
            para = ''
        continue

    if line[0] == '#':
        if para != '':
            doc.add_paragraph(para)
            para = ''
        
        doc.add_heading(line.split(' ')[1])
        continue
    
    if line == '---' or line == '----':
        if para != '':
            doc.add_paragraph(para)
            para = ''
        
        doc.add_page_break()
        continue

    para = para + line

doc.save('sample.docx')